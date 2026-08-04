"""文档处理模块 - 负责文档解析、分块

v0.5 增强: PDF 版面分析(标题/正文识别)、表格识别转 Markdown、内嵌图片 OCR
"""
import threading
from pathlib import Path
from dataclasses import dataclass
from typing import Optional

from app.core.config import DocumentConfig, settings as _default_settings


@dataclass
class DocumentChunk:
    """文档分块"""
    content: str
    metadata: dict
    chunk_id: Optional[str] = None


class _OCRService:
    """图片 OCR 服务 - 基于 RapidOCR(onnxruntime)，模型内置无需下载。

    懒加载 + 线程锁，加载失败自动降级(返回 None)，不影响主流程。
    """

    def __init__(self):
        self._engine = None
        self._lock = threading.Lock()
        self._failed = False

    def ocr_bytes(self, img_bytes: bytes) -> Optional[str]:
        """识别图片中的文字，失败返回 None"""
        engine = self._get_engine()
        if engine is None:
            return None
        try:
            result, _ = engine(img_bytes)
            if not result:
                return None
            lines = [str(item[1]).strip() for item in result if item and item[1]]
            lines = [ln for ln in lines if ln]
            return "\n".join(lines) if lines else None
        except Exception:
            return None

    def _get_engine(self):
        if self._engine is not None or self._failed:
            return self._engine
        with self._lock:
            if self._engine is not None:
                return self._engine
            try:
                from rapidocr_onnxruntime import RapidOCR
                self._engine = RapidOCR()
            except Exception:
                self._failed = True
        return self._engine


# 全局 OCR 服务实例
_ocr_service = _OCRService()


class DocumentParser:
    """文档解析器 - 支持多种格式"""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".md", ".txt", ".csv"}

    def __init__(self, config: Optional[DocumentConfig] = None):
        self.config = config or _default_settings.document

    def parse(self, file_path: str) -> list[DocumentChunk]:
        """解析文档，返回文档块列表"""
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".pdf":
            return self._parse_pdf(path)
        elif ext == ".docx":
            return self._parse_docx(path)
        elif ext == ".xlsx":
            return self._parse_excel(path)
        elif ext == ".md":
            return self._parse_markdown(path)
        elif ext in (".txt", ".csv"):
            return self._parse_text(path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    # ---------------------------------------------------------------- PDF
    def _parse_pdf(self, path: Path) -> list[DocumentChunk]:
        """解析 PDF 文件 - 优先使用增强模式(版面分析+表格+OCR)，失败降级为纯文本抽取"""
        try:
            return self._parse_pdf_layout(path)
        except ImportError:
            pass  # pymupdf 未安装, 降级 pypdf
        except Exception:
            pass  # 版面分析异常, 降级 pypdf 保证可用性
        return self._parse_pdf_plain(path)

    def _parse_pdf_layout(self, path: Path) -> list[DocumentChunk]:
        """PDF 版面分析: 标题/正文识别 + 表格转 Markdown + 内嵌图片 OCR"""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("请安装 pymupdf: pip install pymupdf")

        doc = fitz.open(str(path))
        chunks: list[DocumentChunk] = []
        for page_no in range(len(doc)):
            page = doc[page_no]
            base_meta = {
                "source": str(path),
                "filename": path.name,
                "page": page_no + 1,
                "type": "pdf",
            }
            table_bboxes: list[tuple] = []
            if self.config.table_to_markdown:
                table_chunks, table_bboxes = self._extract_tables(page, base_meta)
                chunks.extend(table_chunks)
            text_chunks = self._extract_text_blocks(page, base_meta, table_bboxes)
            chunks.extend(text_chunks)
            if self.config.ocr_enabled:
                # 收集标题及其 y 坐标, 用于定位图片所属章节
                headings = []
                for c in text_chunks:
                    if c.metadata.get("block_type") != "heading":
                        continue
                    bbox = c.metadata.get("bbox", "")
                    y0 = float(bbox.split(",")[1]) if bbox else 0.0
                    headings.append((y0, c.content))
                chunks.extend(self._extract_images_ocr(doc, page, base_meta, headings))
        doc.close()
        return chunks

    def _extract_text_blocks(self, page, base_meta: dict,
                             table_bboxes: Optional[list[tuple]] = None) -> list[DocumentChunk]:
        """版面分析: 按文本块提取, 依据字号判定标题/正文; 跳过表格区域内的重复文本"""
        try:
            data = page.get_text("dict")
        except Exception:
            return []
        chunks: list[DocumentChunk] = []
        for block in data.get("blocks", []):
            if block.get("type") != 0:  # 仅文本块
                continue
            lines = block.get("lines", [])
            if not lines:
                continue
            bbox = block.get("bbox", (0, 0, 0, 0))
            if table_bboxes and DocumentParser._bbox_in_tables(bbox, table_bboxes):
                continue  # 表格区域内的文本已由表格解析覆盖
            # 拼接块内文本, 记录最大字号(用于标题判定)
            text_parts, max_size = [], 0.0
            for line in lines:
                line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                if line_text.strip():
                    text_parts.append(line_text)
                for span in line.get("spans", []):
                    max_size = max(max_size, span.get("size", 0) or 0)
            text = "\n".join(text_parts).strip()
            if not text:
                continue
            chunks.append(DocumentChunk(
                content=text,
                metadata={
                    **base_meta,
                    "block_type": "heading" if max_size >= self.config.heading_min_size else "body",
                    "font_size": round(max_size, 1),
                    "bbox": ",".join(str(round(v, 1)) for v in bbox),
                }
            ))
        return chunks

    @staticmethod
    def _bbox_in_tables(bbox, table_bboxes: list[tuple]) -> bool:
        """判断文本块 bbox 是否与任一表格 bbox 大面积重叠"""
        x0, y0, x1, y1 = bbox
        area = max(1e-6, (x1 - x0) * (y1 - y0))
        for tb in table_bboxes:
            ix0, iy0 = max(x0, tb[0]), max(y0, tb[1])
            ix1, iy1 = min(x1, tb[2]), min(y1, tb[3])
            if ix1 > ix0 and iy1 > iy0:
                inter = (ix1 - ix0) * (iy1 - iy0)
                if inter / area > 0.5:
                    return True
        return False

    def _extract_tables(self, page, base_meta: dict) -> tuple[list[DocumentChunk], list[tuple]]:
        """表格识别: find_tables 检测页面表格, 转换为 Markdown

        返回 (chunks, table_bboxes)，bbox 用于文本块去重
        """
        try:
            tables = page.find_tables()
        except Exception:
            return [], []
        chunks: list[DocumentChunk] = []
        bboxes: list[tuple] = []
        for tbl in tables.tables:
            rows = tbl.extract()
            if not rows or len(rows) < 2:
                continue
            header = [str(c) if c is not None else "" for c in rows[0]]
            md_lines = [
                "| " + " | ".join(header) + " |",
                "| " + " | ".join("---" for _ in header) + " |",
            ]
            for row in rows[1:]:
                cells = [str(c).replace("\n", " ") if c is not None else "" for c in row]
                md_lines.append("| " + " | ".join(cells) + " |")
            bbox = getattr(tbl, "bbox", (0, 0, 0, 0))
            bboxes.append(tuple(bbox))
            chunks.append(DocumentChunk(
                content="\n".join(md_lines),
                metadata={
                    **base_meta,
                    "block_type": "table",
                    "bbox": ",".join(str(round(v, 1)) for v in bbox),
                }
            ))
        return chunks, bboxes

    def _extract_images_ocr(self, doc, page, base_meta: dict,
                            headings: Optional[list[tuple]] = None) -> list[DocumentChunk]:
        """提取页面内嵌图片并 OCR, 结果作为可检索文本入块

        headings: [(标题y坐标, 标题文本)], 用于定位图片所属章节,
                  拼入内容帮助中文查询命中英文 OCR 结果
        """
        try:
            xrefs = {im[0] for im in page.get_images(full=True)}
        except Exception:
            return []
        chunks: list[DocumentChunk] = []
        ocr_count = 0
        for xref in xrefs:
            if ocr_count >= self.config.ocr_max_images_per_page:
                break
            try:
                rects = page.get_image_rects(xref)
                if not rects:
                    continue
                rect = rects[0]
                if rect.width * rect.height < self.config.ocr_min_area:
                    continue
                info = doc.extract_image(xref)
                if not info or not info.get("image"):
                    continue
                text = _ocr_service.ocr_bytes(info["image"])
                if not text:
                    continue
                ocr_count += 1
                # 取图片上方最近的标题作为章节上下文
                page_context = ""
                if headings:
                    for y0, htext in headings:
                        if y0 < rect.y0:
                            page_context = htext
                prefix = f"[图片内容·{page_context}]" if page_context else "[图片内容]"
                chunks.append(DocumentChunk(
                    content=f"{prefix} {text}",
                    metadata={
                        **base_meta,
                        "block_type": "ocr",
                        "bbox": ",".join(str(round(v, 1)) for v in (rect.x0, rect.y0, rect.x1, rect.y1)),
                    }
                ))
            except Exception:
                continue
        return chunks

    @staticmethod
    def _parse_pdf_plain(path: Path) -> list[DocumentChunk]:
        """降级方案: pypdf 按页纯文本抽取"""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("请安装 pypdf: pip install pypdf")

        reader = PdfReader(str(path))
        chunks = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                chunks.append(DocumentChunk(
                    content=text.strip(),
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "page": i + 1,
                        "type": "pdf",
                    }
                ))
        return chunks

    # ---------------------------------------------------------- 其他格式
    @staticmethod
    def _parse_docx(path: Path) -> list[DocumentChunk]:
        """解析 Word 文档"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = Document(str(path))
        full_text = "\n".join(
            para.text for para in doc.paragraphs if para.text.strip()
        )
        if full_text.strip():
            return [DocumentChunk(
                content=full_text.strip(),
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "type": "docx",
                }
            )]
        return []

    @staticmethod
    def _parse_excel(path: Path) -> list[DocumentChunk]:
        """解析 Excel 文件"""
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ImportError("请安装 openpyxl: pip install openpyxl")

        wb = load_workbook(str(path), read_only=True)
        chunks = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    rows.append(row_text)
            if rows:
                chunks.append(DocumentChunk(
                    content="\n".join(rows),
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "sheet": sheet_name,
                        "type": "xlsx",
                    }
                ))
        wb.close()
        return chunks

    @staticmethod
    def _parse_markdown(path: Path) -> list[DocumentChunk]:
        """解析 Markdown 文件"""
        text = path.read_text(encoding="utf-8")
        if text.strip():
            return [DocumentChunk(
                content=text.strip(),
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "type": "markdown",
                }
            )]
        return []

    @staticmethod
    def _parse_text(path: Path) -> list[DocumentChunk]:
        """解析纯文本文件"""
        text = path.read_text(encoding="utf-8")
        if text.strip():
            return [DocumentChunk(
                content=text.strip(),
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "type": "text",
                }
            )]
        return []


class TextSplitter:
    """文本分块器"""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split(self, chunks: list[DocumentChunk]) -> list[DocumentChunk]:
        """将文档块进一步切分为更小的块"""
        result = []
        for chunk in chunks:
            sub_chunks = self._split_text(chunk.content)
            for i, sub_text in enumerate(sub_chunks):
                result.append(DocumentChunk(
                    content=sub_text,
                    metadata={**chunk.metadata, "chunk_index": i},
                ))
        return result

    def _split_text(self, text: str) -> list[str]:
        """按字符长度分块，保留重叠"""
        if len(text) <= self.chunk_size:
            return [text]

        chunks = []
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - self.chunk_overlap
        return chunks
